# This Python file uses the following encoding: utf-8
#------------------------------------------------------------------------------
# Name:        useful_rid_code.py
# Purpose:     Functions for updated RID analysis.
#
# Author:      James Sample
#
# Created:     13/06/2017
# Copyright:   (c) James Sample and NIVA   
#------------------------------------------------------------------------------

def extract_water_chem(stn_id, par_list, st_dt, end_dt, engine, plot=False,
                       samp_sel=None):
    """ Extracts time series for the specified station and parameters. Returns 
        a dataframe of chemistry and LOD values, a dataframe of duplicates and,
        optionally, an interactive grid plot.

    Args:
        stn_id:   Int. Valid RESA2 STATION_ID
        par_list: List of valid RESA2 PARAMETER_NAMES
        st_dt:    Str. Start date as 'yyyy-mm-dd'
        end_dt:   Str. End date as 'yyyy-mm-dd'
        engine:   SQL-Alchemy 'engine' object already connected to RESA2
        plot:     Bool. Choose whether to return a grid plot as well as the 
                  dataframe
        samp_sel: Int or None. If None, extract all samples for stn_id. If Int,
                  filter samples to include only those associated with this
                  sample_selection_id. Sample selections are defined in
                  RESA2.SAMPLE_SELECTION_DEFINITIONS and RESA2.SAMPLE_SELECTIONS.
                  See:
                  http://nbviewer.jupyter.org/github/JamesSample/rid/blob/master/notebooks/programme_changes_2017-18.ipynb
                  for details

    Returns:
        If plot is False, returns (chem_df, dup_df), otherwise
        returns (chem_df, dup_df, fig_obj)
    """
    import matplotlib.pyplot as plt 
    import datetime as dt
    import pandas as pd
    import seaborn as sn

    # Check stn is valid
    sql = ('SELECT * FROM resa2.stations '
           'WHERE station_id = %s' % stn_id)   

    stn_df = pd.read_sql_query(sql, engine)
    
    assert len(stn_df) == 1, 'Station code not found.' 
    
    # Get station_code
    stn_code = stn_df['station_code'].iloc[0]
    
    # Check pars are valid
    if len(par_list)==1:  
        sql = ("SELECT * FROM resa2.parameter_definitions "
               "WHERE name = '%s'" % par_list[0])
    else:
        sql = ('SELECT * FROM resa2.parameter_definitions '
               'WHERE name in %s' % str(tuple(par_list)))

    par_df = pd.read_sql_query(sql, engine)
    
    assert len(par_df) == len(par_list), 'One or more parameters not found.'
    
    # Get results for ALL pars for sites and period of interest
    # Only get values that are "approved" by lab
    sql = ("SELECT * FROM resa2.water_chemistry_values2 "
           "WHERE sample_id IN ( "
           "  SELECT water_sample_id FROM resa2.water_samples "
           "  WHERE station_id = %s "
           "  AND sample_date <= DATE '%s' "
           "  AND sample_date >= DATE '%s') "
           "AND approved = 'YES'" % (stn_df['station_id'].iloc[0],
                                     end_dt,
                                     st_dt))    

    wc_df = pd.read_sql_query(sql, engine)

    # If desired, subset according to sample selection
    if samp_sel:
        assert isinstance(samp_sel, int), '"samp_sel" must be of type "Int".'
        
        # Get list of samples in this selection
        sql = ("SELECT water_sample_id "
               "FROM resa2.sample_selections "
               "WHERE sample_selection_id = %s" % samp_sel)
        
        samp_sel_df = pd.read_sql_query(sql, engine)

        # Filter wc_df based on these sample_ids
        wc_df = wc_df.query('sample_id in %s' % str(tuple(samp_sel_df['water_sample_id'].values)))
        
    # Get all sample dates for sites and period of interest
    sql = ("SELECT water_sample_id, station_id, sample_date "
           "FROM resa2.water_samples "
           "WHERE station_id = %s "
           "AND sample_date <= DATE '%s' "
           "AND sample_date >= DATE '%s'" % (stn_df['station_id'].iloc[0],
                                             end_dt,
                                             st_dt))    

    samp_df = pd.read_sql_query(sql, engine)
    
    
    # Join in par IDs based on method IDs
    sql = ('SELECT * FROM resa2.wc_parameters_methods')
    meth_par_df = pd.read_sql_query(sql, engine)
    
    wc_df = pd.merge(wc_df, meth_par_df, how='left',
                     left_on='method_id', right_on='wc_method_id')
    
    # Get just the parameters of interest
    wc_df = wc_df.query('wc_parameter_id in %s' % str(tuple(par_df['parameter_id'].values)))
    
    # Join in sample dates
    wc_df = pd.merge(wc_df, samp_df, how='left',
                     left_on='sample_id', right_on='water_sample_id')
    
    # Join in parameter units
    sql = ('SELECT * FROM resa2.parameter_definitions')
    all_par_df = pd.read_sql_query(sql, engine)
    
    wc_df = pd.merge(wc_df, all_par_df, how='left',
                     left_on='wc_parameter_id', right_on='parameter_id')
    
    # Join in station codes
    wc_df = pd.merge(wc_df, stn_df, how='left',
                     left_on='station_id', right_on='station_id')
    
    # Convert units
    wc_df['value'] = wc_df['value'] * wc_df['conversion_factor']
    
    # Extract columns of interest
    wc_df = wc_df[['station_code', 'sample_date', 'name', 'unit', 
                   'value', 'flag1', 'entered_date_x']]

    # Check for duplicates
    dup_df = wc_df[wc_df.duplicated(subset=['station_code',
                                            'sample_date',
                                            'name'], 
                                            keep=False)].sort_values(by=['station_code', 
                                                                         'sample_date', 
                                                                         'name'])
    
    dup_df['station_code'] = dup_df['station_code'].str.decode('windows-1252')
    
    if len(dup_df) > 0:
        print ('    WARNING\n    The database contains duplicated values for some station-'
               'date-parameter combinations.\n    Only the most recent values '
               'will be used, but you should check the repeated values are not '
               'errors.\n    The duplicated entries are returned in a separate '
               'dataframe.\n')
        
        # Choose most recent record for each duplicate
        wc_df.sort_values(by='entered_date_x', inplace=True, ascending=True)

        # Drop duplicates
        wc_df.drop_duplicates(subset=['station_code', 'sample_date', 'name'],
                              keep='last', inplace=True)
        
        # Tidy
        del wc_df['entered_date_x']               
        wc_df.reset_index(inplace=True, drop=True)
        
    if plot:
        # Get number of rows
        if (len(par_list) % 3) == 0:
            rows = int(len(par_list) / 3)
        else:
            rows = int(len(par_list) / 3) + 1
        
        # Setup plot grid
        fig, axes = plt.subplots(nrows=rows, ncols=3,
                                 sharex=False, sharey=False, 
                                 figsize=(12, rows*2.5),
                                 squeeze=False)
        axes = axes.flatten()

        # Loop over series
        idx = 0
        for pid, par in enumerate(par_list):
            # Extract series
            ts = wc_df.query('(station_code==@stn_code) & (name==@par)')

            # Get unit
            unit = ts['unit'].iloc[0]

            # Date index and sort
            ts = ts[['sample_date', 'value']]
            ts.index = ts['sample_date']
            del ts['sample_date']
            ts.sort_index(inplace=True)

            # Plot
            ts.plot(ax=axes[idx], legend=False)
            axes[idx].set_title('%s at %s' % (par, 
                                              stn_code.decode('windows-1252')), 
                                fontsize=16)
            axes[idx].set_xlabel('')
            if unit:
                axes[idx].set_ylabel('%s (%s)' % (par, 
                                                  unit.decode('windows-1252')),
                                     fontsize=14)
            else:
                axes[idx].set_ylabel(par, fontsize=14)
            
            # Set x-limits
            axes[idx].set_xlim([dt.datetime.strptime(st_dt, '%Y-%m-%d').date(),
                                dt.datetime.strptime(end_dt, '%Y-%m-%d').date()]) 
            
            idx += 1

        # Delete unwanted axes
        for idx, ax in enumerate(axes):
            if (idx + 1) > len(par_list):
                fig.delaxes(axes[idx])
        
        # Tidy
        plt.tight_layout()
        
    # Restructure df
    #wc_df['unit'] = wc_df['unit'].str.decode('windows-1252') # Not required for Python 3
    wc_df['par'] = wc_df['name'] + '_' + wc_df['unit'].map(str)
    wc_df['flag'] = wc_df['name'] + '_flag'
    del wc_df['station_code'], wc_df['name'], wc_df['unit']

    # DF of values
    wc_val_df = wc_df[['sample_date', 'par', 'value']]
    wc_val_df = wc_val_df.pivot(index='sample_date', columns='par', values='value')

    # DF of LOD flags
    wc_lod_df = wc_df[['sample_date', 'flag', 'flag1']]
    wc_lod_df = wc_lod_df.pivot(index='sample_date', columns='flag', values='flag1')

    # Join
    wc_df = pd.merge(wc_val_df, wc_lod_df, how='left',
                     left_index=True, right_index=True)
    
    # Sort
    wc_df.sort_index(inplace=True, ascending=True)
    
    if plot:
        return (wc_df, dup_df, fig)
    else:
        return (wc_df, dup_df)

def extract_discharge(stn_id, st_dt, end_dt, engine, plot=False):
    """ Extracts daily flow time series for the selected site. Returns a 
        dataframe and, optionally, a plot object
        
    Args:
        stn_id:    Int. Valid RESA2 STATION_ID
        st_dt:     Str. Start date as 'yyyy-mm-dd'
        end_dt:    Str. End date as 'yyyy-mm-dd'
        engine:    SQL-Alchemy 'engine' object already connected to RESA2
        plot:      Bool. Choose whether to return a grid plot as well as the 
                   dataframe
    
    Returns:
        If plot is False, returns a dataframe of flows, otherwise
        returns a tuple (q_df, figure object)
    """
    import matplotlib.pyplot as plt 
    import datetime as dt
    import pandas as pd
    import seaborn as sn

    # Check stn is valid
    sql = ("SELECT * FROM resa2.stations "
           "WHERE station_id = %s" % stn_id)
    stn_df = pd.read_sql_query(sql, engine)
    
    assert len(stn_df) == 1, 'Error in station code.' 

    # Get station_code
    stn_code = stn_df['station_code'].iloc[0]
    
    # Check a discharge station is defined for this WC station
    sql = ("SELECT * FROM resa2.default_dis_stations "
           "WHERE station_id = '%s'" % stn_df['station_id'].iloc[0])
    dis_df = pd.read_sql_query(sql, engine)
    
    assert len(dis_df) == 1, 'Error identifying discharge station.'
    
    # Get ID for discharge station
    dis_stn_id = dis_df['dis_station_id'].iloc[0]

    # Get catchment areas
    # Discharge station
    sql = ("SELECT area FROM resa2.discharge_stations "
           "WHERE dis_station_id = %s" % dis_stn_id)
    area_df = pd.read_sql_query(sql, engine)    
    dis_area = area_df['area'].iloc[0]

    # Chemistry station
    sql = ("SELECT catchment_area FROM resa2.stations "
           "WHERE station_id = %s" % stn_df['station_id'].iloc[0])
    area_df = pd.read_sql_query(sql, engine)    
    wc_area = area_df['catchment_area'].iloc[0]
    
    # Get the daily discharge data for this station
    sql = ("SELECT xdate, xvalue FROM resa2.discharge_values "
           "WHERE dis_station_id = %s "
           "AND xdate >= DATE '%s' "
           "AND xdate <= DATE '%s'" % (dis_stn_id,
                                       st_dt,
                                       end_dt))
    q_df = pd.read_sql_query(sql, engine)
    q_df.columns = ['date', 'flow_m3/s']
    q_df.index = q_df['date']
    del q_df['date']
    
    # Scale flows by area
    q_df = q_df*wc_area/dis_area
    
    # Convert to daily
    q_df = q_df.resample('D').mean()
    
    # Linear interpolation of NoData gaps
    q_df.interpolate(method='linear', inplace=True)
    
    # Plot
    if plot:
        ax = q_df.plot(legend=False, figsize=(12,5))
        ax.set_xlabel('')
        ax.set_ylabel('Flow (m3/s)', fontsize=16)
        ax.set_title('Discharge at %s' % stn_code.decode('windows-1252'),
                     fontsize=16)
        plt.tight_layout()
        
        return (q_df, ax)
    
    else:    
        return q_df

def adjust_lod_values(wc_df):
    """ Adjusts LOD values according to the OSPAR methodology.
    
    Args:
        wc_df: Dataframe in format returned by rid.extract_water_chem
    
    Returns:
        Modified dataframe
    """
    import pandas as pd
    
    # Get list of chem cols
    cols = wc_df.columns
    par_unit_list = [i for i in cols if i.split('_')[1] != 'flag']

    # loop over cols
    for par_unit in par_unit_list:
        par = par_unit.split('_')[0]
        
        # Get just records for this par
        par_df = wc_df[[par_unit, par+'_flag']].copy()
        par_df.dropna(subset=[par_unit], inplace=True)

        # Get vals
        val_df = par_df[par_unit].values

        # Get LOD flags
        lod_df = par_df[par+'_flag'].values

        # Number of LOD values
        n_lod = (lod_df == '<').sum()

        # Prop <= LOD
        p_lod = (100.*n_lod)/len(lod_df)

        # Adjust ALL values
        ad_val_df = (val_df / 100.)*(100. - p_lod)

        # Update ONLY the LOD values
        val_df[(lod_df=='<')] = ad_val_df[(lod_df=='<')]

        # Update values in wc_df
        par_df[par_unit] = val_df  
        
        # Update values in wc_df
        del wc_df[par_unit]
        wc_df = pd.concat([wc_df, par_df[[par_unit]]], axis=1)

    return wc_df
    
def estimate_loads(stn_id, par_list, year, engine, infer_missing=True, samp_sel=None):
    """ Estimates annual pollutant loads for specified site and year.
        
    Args:
        stn_id:        Int. Valid RESA2 STATION_ID
        par_list:      List of valid RESA2 PARAMETER_NAMES
        year:          Int. Year of interest
        engine:        SQL-Alchemy 'engine' object already connected to 
                       RESA2
        infer_missing: Whether to estimate data for missing years using
                       simple regression (see notebook for full details)
        samp_sel:      Int or None. Passed to extract_water_chem - see 
                       doc-string for details
    
    Returns:
        Dataframe of annual loads. If infer_missing=True, '_Est' columns
        are included for each parameter specifying whether the reported
        value is directly based on observations or inferred from 
        regression. If infer_missing=False, these '_Est' columns are
        omitted.
    """
    import pandas as pd
    import numpy as np
    
    # Dict with factors for unit conversions
    unit_dict = {'SPM':[1.E9, 'tonnes'],     # mg to tonnes
                 'TOC':[1.E9, 'tonnes'],     # mg to tonnes
                 'PO4-P':[1.E12, 'tonnes'],  # ug to tonnes
                 'TOTP':[1.E12, 'tonnes'],   # ug to tonnes
                 'NO3-N':[1.E12, 'tonnes'],  # ug to tonnes
                 'NH4-N':[1.E12, 'tonnes'],  # ug to tonnes
                 'TOTN':[1.E12, 'tonnes'],   # ug to tonnes
                 'SiO2':[1.E9, 'tonnes'],    # mg to tonnes
                 'Ag':[1.E12, 'tonnes'],     # ug to tonnes
                 'As':[1.E12, 'tonnes'],     # ug to tonnes
                 'Pb':[1.E12, 'tonnes'],     # ug to tonnes
                 'Cd':[1.E12, 'tonnes'],     # ug to tonnes
                 'Cu':[1.E12, 'tonnes'],     # ug to tonnes
                 'Zn':[1.E12, 'tonnes'],     # ug to tonnes
                 'Ni':[1.E12, 'tonnes'],     # ug to tonnes
                 'Cr':[1.E12, 'tonnes'],     # ug to tonnes
                 'Hg':[1.E12, 'kg'],         # ng to kg
                 'HCHG':[1.E15, 'tonnes'],   # ng to tonnes
                 'SUMPCB':[1.E15, 'tonnes']} # ng to tonnes
    
    # Get water chem data
    wc_df, dup_df = extract_water_chem(stn_id, par_list, 
                                       '%s-01-01' % year,
                                       '%s-12-31' % year,
                                       engine, plot=False,
                                       samp_sel=samp_sel)

    # Estimate missing values
    if infer_missing:
        wc_df = estimate_missing_data(wc_df, stn_id, par_list, year, 
                                      engine, samp_sel=samp_sel)
           
    if len(wc_df) > 0:
        # Get flow data
        q_df = extract_discharge(stn_id, 
                                 '%s-01-01' % year, 
                                 '%s-12-31' % year,
                                 engine, plot=False)
        q_df['date'] = q_df.index.date
        q_df.reset_index(drop=True, inplace=True)
        
        # Total annual flow (Qr)
        sigma_q = q_df['flow_m3/s'].sum()*60*60*24

        # Dict for data
        data_dict = {}
        
        # Adjust LOD values
        wc_df = adjust_lod_values(wc_df)
        
        # Get list of chem cols
        cols = wc_df.columns
        par_unit_list = [i for i in cols if i.split('_')[1] != 'flag']
        
        # Loop over cols
        for par_unit in par_unit_list:
            # Get whether estimated or observed
            if len(par_unit.split('(')) > 1:
                est = 1
            else:
                est = 0
                
            # Get data for this par and drop any missing
            par_df = wc_df[[par_unit]].dropna()  

            # Tidy 
            par_df.index.name = 'datetime'
            par_df['date'] = par_df.index.date
            par_df.reset_index(inplace=True)        

            # Join on date
            par_df = pd.merge(par_df, q_df, how='left', on='date')

            # Calc intervals t_i
            # Get sample dates
            dates = list(par_df['datetime'].values)

            # Add st and end of year
            dates.insert(0, np.datetime64('%s-01-01' % year))
            dates.append(np.datetime64('%s-01-01' % (year+1))) # Assumes time = midnight
            dates = np.array(dates)

            # Calc differences in seconds between dates and divide by 2
            secs = (np.diff(dates) / np.timedelta64(1, 's')) / 2.

            # Add "before" and "after" intervals to df
            t_wts = secs[1:] + secs[:-1]
            
            # The first sample covers the entire period from the start of the 
            # year to the sampling date (not from halfway through the period 
            # like the rest. The same is true for the last sample to the year 
            # end 
            # => add half a period to start and end
            t_wts[0] = t_wts[0] + secs[0]
            t_wts[-1] = t_wts[-1] + secs[-1]
            
            # Add to df
            par_df['t_i'] = t_wts

            # Estimate loads
            # Denominator
            par_df['Qi_ti'] = par_df['flow_m3/s']*par_df['t_i']

            # Numerator
            par_l = par_unit.split('/')[0]
            par_df[par_l] = 1000*par_df['flow_m3/s']*par_df['t_i']*par_df[par_unit]

            # Flow totals
            sum_df = par_df.sum()
            sigma_Qi_ti = sum_df['Qi_ti'] # Denominator

            # Chem totals
            var_name = par_l.split('_')[0]

            # Get unit and conv factor
            conv_fac = unit_dict[var_name][0]
            unit = unit_dict[var_name][1]
            data_dict[var_name+'_'+unit] = (sigma_q*sum_df[par_l])/(conv_fac*sigma_Qi_ti)
            
            # Add estimated flag
            data_dict[var_name+'_Est'] = est

        # Convert to df
        l_df = pd.DataFrame(data_dict, index=[stn_id])
        
        # Only need "Est" cols if infer_missing=True. Otherwise, just return loads
        if not infer_missing:
            cols = [i for i in l_df.columns if i.split('_')[1] != 'Est']
            l_df = l_df[cols]            

        return l_df

def estimate_missing_data(wc_df, stn_id, par_list, year, engine, samp_sel=None):
    """ If observations are not available for the station-parameter(s)-year
        specified, this function will impute values using approximately the
        methodology coded previously by Tore. See RESA2 procedures
        CALCULATE_RID_STATISTICS and FIXNON109 for Tore's code.
       
    Args:
        wc_df:    Dataframe of real observations for the specified
                  station-par_list-year, as returned by extract_water_chem
        stn_id:   Int. Valid RESA2 station ID
        par_list: List of RESA2 parameters of interest
        year:     Int. Year of interest
        engine:   SQL-Alchemy 'engine' object already connected to RESA2
        samp_sel: Int or None. Passed to extract_water_chem - see 
                  doc-string for details
                       
    Returns:
        Modifed version of wc_df, with estimated values for any missing
        parameters in the specified years
    """
    import pandas as pd
    import numpy as np
    import datetime as dt
    from scipy import stats

    # Work out which pars have no data i.e. difference between cols with data
    # and cols requested originally
    mis_par_list = list(set(par_list) - 
                        set([i.split('_')[0] for i in wc_df.columns]))

    # Are some data missing?
    if len(mis_par_list) > 0: 
        # Get all the data from 1990 to the (current_year - 1)
        wc_ts, dup_ts = extract_water_chem(stn_id, mis_par_list, 
                                           '1990-01-01',
                                           '%s-12-31' % (dt.datetime.now().year - 1),
                                           engine, plot=False, samp_sel=samp_sel)

        # Dict to store estimated values
        est_dict = {}

        # Get list of chem cols with data over longer term
        cols = wc_ts.columns
        par_unit_list = [i for i in cols if i.split('_')[1] != 'flag']  

        # Loop over cols
        for par_unit in par_unit_list:
            # Process col headings
            par = par_unit.split('_')[0]
            
            # Get data for this par
            par_ts = wc_ts[[par_unit, par + '_flag']]

            # Check if data for this par
            if len(par_ts) > 0:
                # Adjust LOD values
                par_ts = adjust_lod_values(par_ts)
            
                # Calculate annual means
                par_ts = par_ts.resample('A').mean()
                par_ts.dropna(how='any', inplace=True)

                # If at least 3 years of data, use Sen's slope
                # to estimate missing values
                if len(par_ts) > 2:
                    # Get first and last years with data
                    first_yr = par_ts.index[0].year
                    last_yr = par_ts.index[-1].year

                    # Determine year on which predictions should be based
                    # (see notebook text for details)
                    if year < first_yr:
                        pred_yr = first_yr
                    elif year > last_yr:
                        pred_yr = last_yr 
                    else:
                        pred_yr = year

                    # Theil slopes with 95% CI on slope
                    slp, incpt, lo, hi = stats.theilslopes(par_ts.values[:,0],
                                                           par_ts.index.year)

                    # If slope significant at p < 0.05, use equn to estimate
                    # unknown year, else use long-term annual median
                    if (lo < 0) and (hi > 0):
                        slp = 0
                        incpt = par_ts.median().values[0]
                   
                    # Estimate conc
                    conc = slp*pred_yr + incpt
                   
                    # Set negative values to 0
                    if conc < 0:
                        conc = 0
                    
                # Otherwise use median of whatever data we have
                else:
                    conc = par_ts.median().values[0]

                # Add to dict
                est_dict[par_unit+'(est)'] = conc
                est_dict[par + '_flag'] = None

        # Create df
        est_df = pd.DataFrame(est_dict, index=[pd.to_datetime('%s-06-30' % year)])

        # Append estimate values to wc_df
        wc_df = pd.concat([wc_df, est_df], axis=1)

    return wc_df

def write_csv_water_chem(stn_df, year, csv_path, engine, samp_sel=None):
    """ Creates  CSV summarising flow and water chemistry data
        for the RID_11 and RID_36 stations. Replaces Tore's 
        NIVA.RESAEXTENSIONS Visual Studio project.
    
    Args:
        stn_df:   Dataframe listing the 47 sites of interest. Must
                  include [station_id, station_code, station_name]
        year:     Int. Year of interest
        csv_path: Filename for CSV
        engine:   SQL-Alchemy 'engine' object already connected to RESA2
        samp_sel: Int or None. Passed to extract_water_chem - see 
                  doc-string for details        
    Returns:
        Dataframe. The CSV is saved.
    """
    import pandas as pd
    import numpy as np

    # Chem pars of interest
    par_list = ['pH', 'KOND', 'TURB860', 'SPM', 'TOC', 'PO4-P', 
                'TOTP', 'NO3-N', 'NH4-N', 'TOTN', 'SiO2', 'Ag', 
                'As', 'Pb', 'Cd', 'Cu', 'Zn', 'Ni', 'Cr', 'Hg']

    # Container for output
    df_list = []

    # Loop over stations
    for stn_id in stn_df['station_id'].values:
        # Get WC data
        wc_df, dup_df  = extract_water_chem(stn_id, par_list, 
                                            '%s-01-01' % year,
                                            '%s-12-31' % year,
                                            engine, plot=False,
                                            samp_sel=samp_sel)
        
        if len(wc_df) == 0:
            print ('No chemistry data found for station ID %s.' % stn_id)
        
        else:
            # Get list of cols of interest for later
            cols = wc_df.columns
            par_unit_list = [i for i in cols if i.split('_')[1] != 'flag']
            par_unit_list.append('Qs_m3/s')

            # Add date col (ignoring time)
            wc_df['date'] = wc_df.index.date

            # Reset index
            wc_df.reset_index(inplace=True)
            
            print ('    Extracting flow data...')

            # Get flow data
            q_df = extract_discharge(stn_id, 
                                     '%s-01-01' % year,
                                     '%s-12-31' % year,
                                     engine, plot=False)

            # Add date col (ignoring time)
            q_df['date'] = q_df.index.date
            q_df.reset_index(drop=True, inplace=True)

            # Join flows to chem
            df = pd.merge(wc_df, q_df, how='left', on='date')
            
            # Set index
            df.index = df['sample_date']

            # Add station id
            df['station_id'] = stn_id

            # Tidy
            df['Qs_m3/s'] = df['flow_m3/s']
            del df['date'], df['flow_m3/s'], df['sample_date']
            df.sort_index(inplace=True)   
            df.reset_index(inplace=True)

            df_list.append(df)

        # Build df
        df = pd.concat(df_list, axis=0)

        # Join stn details
        df = pd.merge(stn_df, df, how='left', on='station_id')

        # Reorder cols and tidy
        st_cols = ['station_id', 'station_code', 'station_name', 'old_rid_group', 
                   'new_rid_group', 'ospar_region', 'sample_date', 'Qs_m3/s']
        unwant_cols = ['nve_vassdrag_nr', 'lat', 'lon', 'utm_north', 'utm_east', 
                       'utm_zone', 'station_type'] 
        par_cols = [i for i in df.columns if i not in (st_cols+unwant_cols)]
        par_cols.sort()

        for col in unwant_cols:
            del df[col]

        df = df[st_cols + par_cols]

        # Write output
        df.to_csv(csv_path, index=False, encoding='utf-8')

    return df

def remove_row(table, row):
    """ Function to remove a row from a Word table. See:
        https://groups.google.com/forum/#!topic/python-docx/aDumlNvK6GM
        
    Args:
        table: Table object
        row:   Row object
    
    Returns:
        None, The table object is modified in-place.
    """
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)
    
def update_cell(row_txt, par_txt, value,
                col_dict, row_dict, tab):
    """ Update a cell in a table specified by row and col names.
    
    Args:
        row_txt:  Str. Label to identify row of interest
        col_txt:  Str. Label to identify col of interest
        value:    New value for cell
        col_dict: Dict indexing columns
        row_dict: Dict indexing rows
        tab:      Table object 
    
    Returns:
        None. The table object is modified in-place.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Get row and col indices
    col = col_dict[par_txt]
    row = row_dict[row_txt]

    # Get cell
    cell = tab.cell(row, col)

    # Modify value
    if isinstance(value, float) or isinstance(value, int):
        cell.text = '%.2f' % value  
    elif isinstance(value, str):
        cell.text = value
    else:
        raise ValueError('Unexpected data type.')

    # Align right
    p = tab.cell(row, col).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def write_word_water_chem_tables(stn_df, year, in_docx, engine, samp_sel=None):
    """ Creates Word tables summarising flow and water chemistry data
        for the RID_11 and RID_36 stations. Replaces Tore's 
        NIVA.RESAEXTENSIONS Visual Studio project.
    
    Args:
        stn_df:   Dataframe listing the 47 sites of interest. Must
                  include [station_id, station_code, station_name]
        year:     Int. Year of interest
        in_docx:  Str. Raw path to Word document. This should be a
                  *COPY* of rid_water_chem_tables_template.docx. Do
                  not use the original template as the files will be 
                  modified
        engine:   SQL-Alchemy 'engine' object already connected to RESA2
        samp_sel: Int or None. Passed to extract_water_chem - see 
                  doc-string for details
        
    Returns:
        None. The specified Word document is modified and saved.
    """
    import pandas as pd
    import numpy as np
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

#    # Chem pars of interest (pre-2017)
#    par_list = ['pH', 'KOND', 'TURB860', 'SPM', 'TOC', 'PO4-P', 
#                'TOTP', 'NO3-N', 'NH4-N', 'TOTN', 'SiO2', 'Ag', 
#                'As', 'Pb', 'Cd', 'Cu', 'Zn', 'Ni', 'Cr', 'Hg']
   
    # Chem pars of interest (2017 onwards; includes DOC, Part. C, TDP
    # and Tot. Part. N) 
    par_list = ['pH', 'KOND', 'TURB860', 'SPM', 'TOC', 'DOC', 'Part. C', 
                'PO4-P', 'TOTP', 'TDP', 'NO3-N', 'NH4-N', 'TOTN', 
                'Tot. Part. N', 'SiO2', 'Ag', 'As', 'Pb', 'Cd', 'Cu', 
                'Zn', 'Ni', 'Cr', 'Hg']

    # Open the document
    doc = Document(in_docx)

    # Set styles for 'Normal' template in this doc
    style = doc.styles['Normal']

    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8)

    p_format = style.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    # Loop over tables
    for tab_id, tab in enumerate(doc.tables):
        # Get station name
        stn_name = tab.cell(0, 0).text

        print ('Processing:', stn_name)

        # Get station ID
        stn_id = stn_df.query('station_name == @stn_name')['station_id'].values[0]

        print ('    Extracting water chemistry data...')

        # Get WC data
        wc_df, dup_df  = extract_water_chem(stn_id, par_list, 
                                            '%s-01-01' % year,
                                            '%s-12-31' % year,
                                            engine, plot=False,
                                            samp_sel=samp_sel)

        # Get list of cols of interest for later
        cols = wc_df.columns
        par_unit_list = [i for i in cols if i.split('_')[1] != 'flag']
        par_unit_list.append('Qs_m3/s')

        # Add date col (ignoring time)
        wc_df['date'] = wc_df.index.date

        # Reset index
        wc_df.reset_index(inplace=True)

        print ('    Extracting flow data...')

        # Get flow data
        q_df = extract_discharge(stn_id, 
                                 '%s-01-01' % year,
                                 '%s-12-31' % year,
                                 engine, plot=False)

        # Add date col (ignoring time)
        q_df['date'] = q_df.index.date
        q_df.reset_index(drop=True, inplace=True)

        # Join flows to chem
        df = pd.merge(wc_df, q_df, how='left', on='date')

        # Set index
        df.index = df['sample_date']

        # Tidy
        df['Qs_m3/s'] = df['flow_m3/s']
        del df['date'], df['flow_m3/s'], df['sample_date']
        df.sort_index(inplace=True)

        print ('    Writing sample dates...')

        # Write sample dates to first col
        dates = df.index.values

        for idx, dt in enumerate(dates):
            # Get cell
            cell = tab.cell(idx+3, 0) # Skip 3 rows of header

            # Modify value
            #cell.text = pd.to_datetime(str(dt)).strftime('%d.%m.%Y %H:%M') # Used pre-2017
            cell.text = pd.to_datetime(str(dt)).strftime('%d.%m.%Y')        # Used 2017 onwards to save space

            # Align right
            p = tab.cell(idx+3, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 

        print ('    Deleting empty rows...')

        # Delete empty rows (each blank table has space for 20 samples, but 
        # usually there are fewer)
        # Work out how many rows to delete
        first_blank = len(dates)+3
        n_blank = 23 - len(dates) - 3
        blank_ids = [first_blank,]*n_blank # Delete first blank row n times 

        for idx in blank_ids:
            # Delete rows
            row = tab.rows[idx]
            remove_row(tab, row)     

        print ('    Writing data values...')

        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(1)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        # Loop over df cols
        for idx, df_row in df.iterrows():
            # Get the date
            #dt_tm = idx.strftime('%d.%m.%Y %H:%M') # Used pre-2017
            dt_tm = idx.strftime('%d.%m.%Y')        # Used 2017 onwards to save space

            # Loop over variables
            for par_unit in par_unit_list:
                # Get just the par
                par = par_unit.split('_')[0]
                
                # Deal with flags
                if par == 'Qs':
                    value = df_row[par_unit]
                elif df_row[par + '_flag'] == '<':
                    value = '<' + str(df_row[par_unit])
                else:
                    value = df_row[par_unit]
                
                # Print NaN as blank
                if pd.isnull(value):
                    value = ''
                    
                # Update the table
                update_cell(dt_tm, par, value,
                            col_dict, row_dict, tab)

        print ('    Writing summary statistics...')

        # Add flag col (all None) for Qs
        df['Qs_flag'] = None

        # Loop over cols
        for df_col in par_unit_list:
            # Get just the par
            par = df_col.split('_')[0]
            
            # Get just data for this par
            samp_df = df[[df_col, par+'_flag']].dropna(subset=[df_col,])

            # Calc statistics
            # 1. Lower av. - assume all LOD values are 0
            # Get vals
            val_df = samp_df[df_col].values.copy()

            # Get LOD flags
            lod_df = samp_df[par+'_flag'].values

            # Update ONLY the LOD values
            val_df[(lod_df=='<')] = 0

            # Average
            lo_av = val_df.mean()
            update_cell('Lower avg.', par, lo_av,
                        col_dict, row_dict, tab)   

            # 2. Upper av. - assume all LOD values are LOD
            up_av = samp_df[df_col].mean()
            update_cell('Upper avg..', par, up_av,
                        col_dict, row_dict, tab)

            # 3. Min
            par_min = samp_df[df_col].min()
            update_cell('Minimum', par, par_min,
                        col_dict, row_dict, tab)

            # 4. Max
            par_max = samp_df[df_col].max()
            update_cell('Maximum', par, par_max,
                        col_dict, row_dict, tab)

            # 5. More than 70% above LOD?
            lod_df = samp_df[par+'_flag'].values
            pct_lod = (100.*(lod_df=='<').sum())/len(lod_df) # % at or below LOD
            pct_lod = 100 - pct_lod                          # % above LOD
            if pct_lod > 70:
                lod_gt_70 = 'yes'
            else:
                lod_gt_70 = 'no'
            update_cell('More than 70% >LOD', par, lod_gt_70,
                        col_dict, row_dict, tab)

            # 6. n samples
            n_samps = len(samp_df)
            update_cell('n', par, n_samps,
                        col_dict, row_dict, tab)

            # 7. Std. Dev.
            par_std = samp_df[df_col].std() 
            update_cell('St.dev', par, par_std,
                        col_dict, row_dict, tab)

        print ('    Done.')

        # Save after each table
        doc.save(in_docx)

    print ('Finished.')
    
def write_word_loads_table(stn_df, loads_csv, in_docx, engine):
    """ Creates Word table summarising annual loads for the 155 main
        RID sites (RID_11 + RID_36 + RID_108). Note: the year is 
        automatically extracted from the file name of loads_csv, which
        is created as an output from the "loads notebook".
    
    Args:
        stn_df:    Dataframe listing the 155 sites of interest. Must
                   include [station_id, station_code, station_name]
        loads_csv: Str. Raw path to CSV summarising loads, as generated 
                   by "loads notebook"
        in_docx:   Str. Raw path to Word document. This should be a
                   *COPY* of rid_loads_by_river_template.docx. Do not
                   use the original template as the file will be 
                   modified
        engine:    SQL-Alchemy 'engine' object already connected to 
                   RESA2
        
    Returns:
        None. The specified Word document is modified and saved.
    """
    import pandas as pd
    import numpy as np
    import os
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Get year from loads_csv
    year = int(os.path.split(loads_csv)[1].split('_')[-1][:-4])

    # Read loads data
    lds_df = pd.read_csv(loads_csv, index_col=0)
    
    # Chem pars of interest
    par_list = ['SPM', 'TOC', 'PO4-P', 'TOTP', 'NO3-N', 'NH4-N', 
                'TOTN', 'SiO2', 'Ag', 'As', 'Pb', 'Cd', 'Cu',
                'Zn', 'Ni', 'Cr', 'Hg']

    # Open the Word document
    doc = Document(in_docx)

    # Set styles for 'Normal' template in this doc
    style = doc.styles['Normal']

    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8)

    p_format = style.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    # Get table obj
    tab = doc.tables[0]

    # Extract text to index rows
    row_dict = {}
    for idx, cell in enumerate(tab.column_cells(0)):
        for paragraph in cell.paragraphs:
            row_dict[paragraph.text] = idx 

    # Extract text to index cols
    col_dict = {}
    for idx, cell in enumerate(tab.row_cells(0)):
        for paragraph in cell.paragraphs:
            col_dict[paragraph.text] = idx 

    # Loop over sites
    print ('Processing:')
    for stn_id in stn_df['station_id']:
        # Get name and code
        name = stn_df.query('station_id == @stn_id')['station_name'].values[0]
        code = stn_df.query('station_id == @stn_id')['station_code'].values[0]

        print ('    %s (%s)...' % (name, code))
        
        # Allow for sites with the same name
        if name in [u'Børselva', u'Oselva']:
            name = '%s (%s)' % (name, code)

        # Get flow data
        q_df = extract_discharge(stn_id, 
                                 '%s-01-01' % year,
                                 '%s-12-31' % year,
                                 engine, plot=False)

        # Average daily flow vol in 1000s m3/day
        q_av = q_df.mean()['flow_m3/s']
        v_av = q_av*24*60*60 / 1000

        # Update the table with flow
        update_cell(name, 'Flow rate', v_av,
                    col_dict, row_dict, tab)

        # Loop over chem pars
        for par in par_list:
            # Get col for loads df
            if par == 'Hg':
                par_l = 'Hg_kg'
            else:
                par_l = par + '_tonnes'

            # Get load value
            load = lds_df.ix[stn_id, par_l]

            # Update table
            update_cell(name, par, load,
                        col_dict, row_dict, tab)

    # Save after each table
    doc.save(in_docx)

    print( 'Finished.')

def write_word_overall_table(mon_df, umon_df, in_docx):
    """ Creates Word tables summarising overall loads for the RID project.
        This code is VERY messy and needs tidying up!
    
    Args:
        mon_df:  Dataframe of monitoring results, created in section 4
                 of word_data_tables.ipynb
        umon_df: Dataframe of modelling results, created in section 4
                 of word_data_tables.ipynb
        in_docx: Str. Raw path to Word document. This should be a
                 *COPY* of rid_loads_overall_summary_template.docx. Do
                 not use the original template as the files will be 
                 modified
        
    Returns:
        None. The specified Word document is modified and saved.
    """
    import pandas as pd
    import numpy as np
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Dicts mapping df cols to Word cols
    reg_dict = {'TOTAL NORWAY':'NORWAY',
                'SKAGGERAK':'SKAGERAK',
                'NORTH SEA':'NORTH SEA',
                'NORWEGIAN SEA':'NORWEGIAN SEA2',
                'LOFOTEN/BARENTS SEA':'LOFOTEN-BARENTS SEA'}

    grp_dict = {'rid_11':'Main Rivers (11)',
                'rid_36':'Tributary Rivers (36)',
                'rid_108':'Tributary Rivers (108)'}

    par_dict = {'Flow rate':'flow', 
                'SPM':'S.P.M.',
                'TOC':'TOC',
                'PO4-P':'po4',
                'TOTP':'p',
                'NO3-N':'no3', 
                'NH4-N':'nh4',
                'TOTN':'n',
                'SiO2':None,
                'Ag':None,
                'As':'As',
                'Pb':'Pb',
                'Cd':'Cd', 
                'Cu':'Cu',
                'Zn':'Zn',
                'Ni':'Ni',
                'Cr':'Cr',
                'Hg':'Hg'}

    typ_dict = {'Sewage Effluents':'sew',
                'Industrial Effluents':'ind',
                'Fish Farming':'fish'}

    # Chem pars of interest
    par_list = ['Flow rate', 'SPM', 'TOC', 'PO4-P', 'TOTP', 'NO3-N', 
                'NH4-N', 'TOTN', 'SiO2', 'Ag', 'As', 'Pb', 'Cd', 
                'Cu', 'Zn', 'Ni', 'Cr', 'Hg']

    # Open the document
    doc = Document(in_docx)

    # Set styles for 'Normal' template in this doc
    style = doc.styles['Normal']

    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8)

    p_format = style.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    # 1. Fill-in observed data
    # Loop over tables
    for tab_id, tab in enumerate(doc.tables):
        # Get region name
        reg = tab.cell(0, 0).text
        reg_df = reg_dict[reg]

        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        # Loop over pars
        for par in par_list:
            # Get vals for RID groups
            df = mon_df.loc[reg_df][[par]]

            # Write total rows
            tot = round(df.values.sum(), 0)

            # Handle NaN
            if np.isnan(tot):
                tot = '0'
            else:
                tot = str(int(tot))

            # Update doc
            update_cell('Total Riverine Inputs', par, tot,
                        col_dict, row_dict, tab)

            # Loop over df
            for idx, row in df.iterrows():
                # Get value
                val = round(row[par], 0)

                # Handle NaN
                if np.isnan(val):
                    val = '0'
                else:
                    val = str(int(val))

                # Write values
                update_cell(grp_dict[idx], par, val,
                            col_dict, row_dict, tab)

    # 2. Fill-in "direct" discharges
    # Loop over tables
    for tab_id, tab in enumerate(doc.tables):
        # Get region name
        reg = tab.cell(0, 0).text
        reg_df = reg_dict[reg]

        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        # Loop over pars
        for par in par_list:
            # Only process relevant rows
            if par_dict[par]:
                # Flow
                if par == 'Flow rate':
                    pass

                else:
                    # Get vals for sew, ind and aqu
                    try:
                        sew = umon_df.loc[reg_df]['sew_'+par_dict[par]]
                    except KeyError:
                        sew = np.nan

                    try:
                        ind = umon_df.loc[reg_df]['ind_'+par_dict[par]]
                    except KeyError:
                        ind = np.nan

                    try:
                        aqu = umon_df.loc[reg_df]['fish_'+par_dict[par]]
                    except KeyError:
                        aqu = np.nan

                    # Get totals
                    tot = np.array([sew, ind, aqu])
                    tot = np.nansum(tot)
                    if tot == 0:
                        tot = ''
                    else:
                        tot = str(int(round(tot)))                    

                    # Format numbers
                    if np.isnan(sew):
                        sew = ''
                    else:
                        sew = str(int(round(sew)))

                    if np.isnan(ind):
                        ind = ''
                    else:
                        ind = str(int(round(ind)))

                    if np.isnan(aqu):
                        aqu = ''
                    else:
                        aqu = str(int(round(aqu)))

                    # Update values
                    update_cell('Sewage Effluents', par, sew,
                                col_dict, row_dict, tab)

                    update_cell('Industrial Effluents', par, ind,
                                col_dict, row_dict, tab)

                    update_cell('Fish Farming', par, aqu,
                                col_dict, row_dict, tab)

                    update_cell('Total Direct Inputs', par, tot,
                                col_dict, row_dict, tab)

    # 3. Process "diffuse"
    for tab_id, tab in enumerate(doc.tables):
        # Get region name
        reg = tab.cell(0, 0).text
        reg_df = reg_dict[reg]

        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        # Loop over pars
        for par in ['Flow rate', 'PO4-P', 'TOTP', 'NO3-N', 'NH4-N', 'TOTN']:
            if par == 'Flow rate':
                val = str(int(round(umon_df.loc[reg_df][par_dict[par]], 0)))
            else:
                val = str(int(round(umon_df.loc[reg_df]['diff_'+par_dict[par]], 0)))

            # Write values. 'Unmonitored values' is in the Word template,
            # but the text is coloured grey so that it's invisible
            update_cell('Unmonitored values', par, val,
                        col_dict, row_dict, tab)

    # 4. Process region totals
    for tab_id, tab in enumerate(doc.tables):
        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        for idx, par in enumerate(par_list):
            riv = tab.cell(6, idx+1).text
            direc = tab.cell(11, idx+1).text
            umon = tab.cell(13, idx+1).text

            # Get total
            tot = 0
            for val in [riv, direc, umon]:
                if val == '':
                    pass
                else:
                    tot += int(val)

            tot = str(int(tot))

            # Write values. 'Unmonitored values' is in the Word template,
            # but the text is coloured grey so that it's invisible
            update_cell('Region values', par, tot,
                        col_dict, row_dict, tab)

    # Save after each table
    doc.save(in_docx)
    
    print ('Finished.')

def write_word_overall_table_2017_20(mon_df, umon_df, in_docx):
    """ Creates Word tables summarising overall loads for the RID project.
        This code is VERY messy and needs tidying up!
    
    Args:
        mon_df:  Dataframe of monitoring results, created in section 4
                 of word_data_tables.ipynb
        umon_df: Dataframe of modelling results, created in section 4
                 of word_data_tables.ipynb
        in_docx: Str. Raw path to Word document. This should be a
                 *COPY* of rid_loads_overall_summary_template.docx. Do
                 not use the original template as the files will be 
                 modified
        
    Returns:
        None. The specified Word document is modified and saved.
    """
    import pandas as pd
    import numpy as np
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Dicts mapping df cols to Word cols
    reg_dict = {'TOTAL NORWAY':'NORWAY',
                'SKAGGERAK':'SKAGERAK',
                'NORTH SEA':'NORTH SEA',
                'NORWEGIAN SEA':'NORWEGIAN SEA2',
                'LOFOTEN/BARENTS SEA':'LOFOTEN-BARENTS SEA'}

    grp_dict = {'rid_20':'Main Rivers (20)',
                'rid_135':'Other Rivers (135)'}

    par_dict = {'Flow rate':'flow', 
                'SPM':'S.P.M.',
                'TOC':'TOC',
                'PO4-P':'po4',
                'TOTP':'p',
                'NO3-N':'no3', 
                'NH4-N':'nh4',
                'TOTN':'n',
                'SiO2':None,
                'Ag':None,
                'As':'As',
                'Pb':'Pb',
                'Cd':'Cd', 
                'Cu':'Cu',
                'Zn':'Zn',
                'Ni':'Ni',
                'Cr':'Cr',
                'Hg':'Hg'}

    typ_dict = {'Sewage Effluents':'sew',
                'Industrial Effluents':'ind',
                'Fish Farming':'fish'}

    # Chem pars of interest
    par_list = ['Flow rate', 'SPM', 'TOC', 'PO4-P', 'TOTP', 'NO3-N', 
                'NH4-N', 'TOTN', 'SiO2', 'Ag', 'As', 'Pb', 'Cd', 
                'Cu', 'Zn', 'Ni', 'Cr', 'Hg']

    # Open the document
    doc = Document(in_docx)

    # Set styles for 'Normal' template in this doc
    style = doc.styles['Normal']

    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8)

    p_format = style.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    # 1. Fill-in observed data
    # Loop over tables
    for tab_id, tab in enumerate(doc.tables):
        # Get region name
        reg = tab.cell(0, 0).text
        reg_df = reg_dict[reg]

        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        # Loop over pars
        for par in par_list:
            # Get vals for RID groups
            df = mon_df.loc[reg_df][[par]]

            # Write total rows
            tot = round(df.values.sum(), 0)

            # Handle NaN
            if np.isnan(tot):
                tot = '0'
            else:
                tot = str(int(tot))

            # Update doc
            update_cell('Total Riverine Inputs', par, tot,
                        col_dict, row_dict, tab)

            # Loop over df
            for idx, row in df.iterrows():
                # Get value
                val = round(row[par], 0)

                # Handle NaN
                if np.isnan(val):
                    val = '0'
                else:
                    val = str(int(val))

                # Write values
                update_cell(grp_dict[idx], par, val,
                            col_dict, row_dict, tab)

    # 2. Fill-in "direct" discharges
    # Loop over tables
    for tab_id, tab in enumerate(doc.tables):
        # Get region name
        reg = tab.cell(0, 0).text
        reg_df = reg_dict[reg]

        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        # Loop over pars
        for par in par_list:
            # Only process relevant rows
            if par_dict[par]:
                # Flow
                if par == 'Flow rate':
                    pass

                else:
                    # Get vals for sew, ind and aqu
                    try:
                        sew = umon_df.loc[reg_df]['sew_'+par_dict[par]]
                    except KeyError:
                        sew = np.nan

                    try:
                        ind = umon_df.loc[reg_df]['ind_'+par_dict[par]]
                    except KeyError:
                        ind = np.nan

                    try:
                        aqu = umon_df.loc[reg_df]['fish_'+par_dict[par]]
                    except KeyError:
                        aqu = np.nan

                    # Get totals
                    tot = np.array([sew, ind, aqu])
                    tot = np.nansum(tot)
                    if tot == 0:
                        tot = ''
                    else:
                        tot = str(int(round(tot)))                    

                    # Format numbers
                    if np.isnan(sew):
                        sew = ''
                    else:
                        sew = str(int(round(sew)))

                    if np.isnan(ind):
                        ind = ''
                    else:
                        ind = str(int(round(ind)))

                    if np.isnan(aqu):
                        aqu = ''
                    else:
                        aqu = str(int(round(aqu)))

                    # Update values
                    update_cell('Sewage Effluents', par, sew,
                                col_dict, row_dict, tab)

                    update_cell('Industrial Effluents', par, ind,
                                col_dict, row_dict, tab)

                    update_cell('Fish Farming', par, aqu,
                                col_dict, row_dict, tab)

                    update_cell('Total Direct Inputs', par, tot,
                                col_dict, row_dict, tab)

    # 3. Process "diffuse"
    for tab_id, tab in enumerate(doc.tables):
        # Get region name
        reg = tab.cell(0, 0).text
        reg_df = reg_dict[reg]

        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        # Loop over pars
        for par in ['Flow rate', 'PO4-P', 'TOTP', 'NO3-N', 'NH4-N', 'TOTN']:
            if par == 'Flow rate':
                val = str(int(round(umon_df.loc[reg_df][par_dict[par]], 0)))
            else:
                val = str(int(round(umon_df.loc[reg_df]['diff_'+par_dict[par]], 0)))

            # Write values. 'Unmonitored values' is in the Word template,
            # but the text is coloured grey so that it's invisible
            update_cell('Unmonitored values', par, val,
                        col_dict, row_dict, tab)

    # 4. Process region totals
    for tab_id, tab in enumerate(doc.tables):
        # Extract text to index rows
        row_dict = {}
        for idx, cell in enumerate(tab.column_cells(0)):
            for paragraph in cell.paragraphs:
                row_dict[paragraph.text] = idx 

        # Extract text to index cols
        col_dict = {}
        for idx, cell in enumerate(tab.row_cells(0)):
            for paragraph in cell.paragraphs:
                col_dict[paragraph.text] = idx 

        for idx, par in enumerate(par_list):
            riv = tab.cell(5, idx+1).text
            direc = tab.cell(10, idx+1).text
            umon = tab.cell(12, idx+1).text

            # Get total
            tot = 0
            for val in [riv, direc, umon]:
                if val == '':
                    pass
                else:
                    tot += int(val)

            tot = str(int(tot))

            # Write values. 'Unmonitored values' is in the Word template,
            # but the text is coloured grey so that it's invisible
            update_cell('Region values', par, tot,
                        col_dict, row_dict, tab)

    # Save after each table
    doc.save(in_docx)
    
    print ('Finished.')
        
def identify_point_in_polygon(pt_df, 
                              poly_shp, 
                              pt_col='ANLEGGSNR', 
                              poly_col='VASSDRAGNR', 
                              lat_col='Lat',
                              lon_col='Lon'):
    """ Performs spatial join to identify containing polygon IDs for points
        with lat/lon co-ordinates.
        
    Args:
        pt_df:    Dataframe of point locations. Must include a unique site 
                  ID and cols containing lat and lon in WGS84 decimal degrees
        poly_shp: Str. Raw path to polygon shapefile. Also in WGS84 geographic
                  co-ordinates
        pt_col:   Str. Name of col with unique point IDs
        poly_col: Str. Name of col with unique polygon IDs
        lat_col:  Str. Name of lat col for points
        lon_col:  Str. Name of lon col for points
    
    Returns:
        Dataframe with "poly_col" column added specifying polygon ID.
    """
    import pandas as pd
    import geopandas as gpd
    import geopandas.tools
    import pyproj
    import numpy as np
    from shapely.geometry import Point
    
    # Get just the spatial info and site IDs
    pt_df2 = pt_df[[pt_col, lat_col, lon_col]].copy()

    # Drop any rows without lat/lon from df
    # NOTE: Ideally, these should have been corrected above!
    if pt_df2.isnull().values.any():
        print ('WARNING: Not all sites have complete co-ordinate information. '
               'These rows will be dropped.')
        pt_df2.dropna(how='any', inplace=True)
        
    # Reset index (otherwise GPD join doesn't work)
    pt_df2.reset_index(inplace=True, drop=True) 

    # Create the geometry column from point coordinates
    pt_df2['geometry'] = pt_df2.apply(lambda row: Point(row[lon_col], 
                                                        row[lat_col]), axis=1)

    # Convert to GeoDataFrame
    pt_df2 = gpd.GeoDataFrame(pt_df2, geometry='geometry')
    del pt_df2[lat_col], pt_df2[lon_col]

    # Set coordinate system as WGS84
    pt_df2.crs = {"init": "epsg:4326"}

    # Load Regine catchment shapefile (projected to WGS84)
    reg_gdf = gpd.GeoDataFrame.from_file(poly_shp)

    # Get cols of interest
    reg_gdf = reg_gdf[[poly_col, 'geometry']]

    # Some vassdragsnummers are duplicated
    reg_gdf.drop_duplicates(subset=[poly_col], inplace=True)
    
    # Spatial join
    join_gdf = gpd.tools.sjoin(pt_df2, reg_gdf,
                               how='left', op='within')

    # Join output back to original data table
    reg_gdf = join_gdf[[pt_col, poly_col]]
    res_df = pd.merge(pt_df, reg_gdf, how='left', on=pt_col)

    return res_df

def utm_to_wgs84_dd(utm_df, zone, east, north):
    """ Converts UTM co-ordinates to WGS84 decimal degrees.
    
    Args:
        utm_df: Dataframe containing UTM co-ords
        zone:   Str. Column defining UTM zone
        east:   Str. Column defining UTM Easting
        north:  Str. Column defining UTM Northing
        
    Returns:
        Copy of utm_df with 'lat' and 'lon' columns added.
    """
    import pandas as pd
    import numpy as np
    import pyproj        
    
    # Copy utm_df
    df = utm_df.copy()
    
    # Containers for data
    lats = []
    lons = []

    # Loop over df
    for idx, row in df.iterrows():
        # Only convert if UTM co-ords are available
        if pd.isnull(row[zone]):
            lats.append(np.nan)
            lons.append(np.nan)    
        else:    
            # Build projection
            p = pyproj.Proj(proj='utm', zone=row[zone], ellps='WGS84')

            # Convert
            lon, lat = p(row[east], row[north], inverse=True)
            lats.append(lat)
            lons.append(lon)

    # Add to df
    df['lat'] = lats
    df['lon'] = lons
    
    return df

def write_teotil_point_source_input_file(out_fold, ts_id,
                                         year, typ, engine):
    """ Writes a text file for input to TEOTIL based on data
        in the RESA2.RID_KILDER_TIMESERIES table.
    
    Args:
        out_fold: Str. Path to folder for output .txt file
        ts_id:    Int. Valid TIMESERIES_ID defined in
                  RESA2.RID_TIMESERIES_DEFINITIONS
        year:     Int. Year of interest
        typ:      One of ['RENSEANLEGG', 'SPREDT', 'INDUSTRI',
                  'AKVAKULTUR']
        engine:   SQL-Alchemy 'engine' object already connected to 
                  RESA2
    
    Returns:
        None. The text file is written to the specified folder.
    """
    import pandas as pd
    import os
    
    assert typ in ['RENSEANLEGG', 'SPREDT', 
                   'INDUSTRI', 'AKVAKULTUR'], '"typ" not valid.'
    
    # Get data from RESA2.RID_KILDER_TIMESERIES
    sql = ("SELECT id, regine, parameter_id, xvalue "
           "FROM resa2.rid_kilder_timeseries "
           "WHERE timeseries_id = %s "
           "AND type = '%s' "
           "AND year = %s "
           "ORDER BY id" % (ts_id, typ, year))

    val_df = pd.read_sql_query(sql, engine)

    # Get data for pars
    sql = ("SELECT DISTINCT out_pid, name, unit "
           "FROM resa2.rid_punktkilder_outpar_def")

    par_df = pd.read_sql_query(sql, engine)

    # Pivot to 'wide' format; pad with zeros; tidy
    val_df.set_index(['id', 'regine', 'parameter_id'], inplace=True)
    val_df = val_df.unstack(level=-1, fill_value=0)
    val_df.columns = val_df.columns.droplevel(0)
    val_df.reset_index(inplace=True)
    val_df.columns.name = ''

    # Add year col
    val_df['year'] = year

    # Add blank cols for any missing pars
    for pid in par_df['out_pid'].values:
        if not pid in val_df.columns:
            val_df[pid] = 0

    # Reorder cols
    val_df = val_df[['id', 'regine', 'year'] + list(par_df['out_pid'].values)]

    # Rename cols
    val_df.columns = ['ID', 'REGINE', 'YEAR'] + list(par_df['name'].values)

    # Build custom headers for TEOTIL
    # 1st row
    row1_txt = '!Dette er en automatisk generert fil. Dato:%s' % pd.to_datetime('now')
    row1 = [row1_txt] + (len(val_df.columns) - 1)*['',]

    # 2nd row
    row2 = (['!ID', 'REGINE', 'YEAR'] + 
            ['%s(%s)' % (par, par_df['unit'].values[idx]) 
             for idx, par in enumerate(par_df['name'].values)])

    # 3rd row
    row3 = val_df.columns

    # Assign header as multi-index
    val_df.columns = pd.MultiIndex.from_tuples(zip(row1, row2, row3))
    
    # Write output
    out_path = os.path.join(out_fold, '%s.txt' % typ)
    val_df.to_csv(out_path, sep=';', index=False, encoding='utf-8')
    
def write_teotil_discharge_input_file(out_fold, year, engine):
    """ Writes a text file for input to TEOTIL based on data
        in the RESA2.DISCHARGE_VALUES table.
    
    Args:
        out_fold: Str. Path to folder for output .txt file
        year:     Int. Year of interest
        engine:   SQL-Alchemy 'engine' object already connected to 
                  RESA2
    
    Returns:
        None. The text file is written to the specified folder.
    """
    import pandas as pd
    import os

    # Get NVE stn IDs
    sql = ("SELECT dis_station_id, TO_NUMBER(nve_serienummer) as Vassdrag "
           "FROM resa2.discharge_stations "
           "WHERE dis_station_name LIKE 'NVE Modellert%'")

    nve_stn_df = pd.read_sql_query(sql, engine)
    nve_stn_df.index = nve_stn_df['dis_station_id']
    del nve_stn_df['dis_station_id']

    # Get avg. annual values for NVE stns
    sql = ("SELECT dis_station_id, ROUND(AVG(xvalue), 6) AS QAr "
           "FROM resa2.discharge_values "
           "WHERE dis_station_id in ( "
           "  SELECT dis_station_id "
           "  FROM resa2.discharge_stations "
           "  WHERE dis_station_name LIKE 'NVE Modellert%%') "
           "AND TO_CHAR(xdate, 'YYYY') = %s "
           "GROUP BY dis_station_id "
           "ORDER BY dis_station_id" % year)

    an_avg_df = pd.read_sql_query(sql, engine)
    an_avg_df.index = an_avg_df['dis_station_id']
    del an_avg_df['dis_station_id']

    ## Get long-term avg. values for NVE stns
    sql = ("SELECT dis_station_id, ROUND(AVG(xvalue), 6) AS Qsnitt "
           "FROM resa2.discharge_values "
           "WHERE dis_station_id in ( "
           "  SELECT dis_station_id "
           "  FROM resa2.discharge_stations "
           "  WHERE dis_station_name LIKE 'NVE Modellert%') "
           "AND TO_CHAR(xdate, 'YYYY') BETWEEN 1990 AND 2008 "
           "GROUP BY dis_station_id "
           "ORDER BY dis_station_id")

    lt_avg_df = pd.read_sql_query(sql, engine)
    lt_avg_df.index = lt_avg_df['dis_station_id']
    del lt_avg_df['dis_station_id']

    # Get monthly avg. values for NVE stns
    sql = ("SELECT dis_station_id, TO_CHAR(xdate,'MM') AS month, "
           "  ROUND(AVG(xvalue), 6) as avg "
           "FROM resa2.discharge_values "
           "WHERE dis_station_id in ( "
           "  SELECT dis_station_id "
           "  FROM resa2.discharge_stations "
           "  WHERE dis_station_name LIKE 'NVE Modellert%%') "
           "AND TO_CHAR(xdate, 'YYYY') = %s "
           "GROUP BY dis_station_id, TO_CHAR(xdate,'MM') "
           "ORDER BY dis_station_id, TO_CHAR(xdate,'MM')" % year)

    mon_avg_df = pd.read_sql_query(sql, engine)
    mon_avg_df.set_index(['dis_station_id', 'month'], inplace=True)
    mon_avg_df = mon_avg_df.unstack(level=-1)
    mon_avg_df.columns = mon_avg_df.columns.droplevel(0)
    mon_avg_df.columns = ['Q%s' % i for i in range(1, 13)]
    mon_avg_df.columns.name = ''

    # Combine
    q_df = pd.concat([nve_stn_df, an_avg_df, 
                      lt_avg_df, mon_avg_df], axis=1)

    # Tidy
    q_df.reset_index(inplace=True, drop=True)
    q_df.sort_values(by='vassdrag', ascending=True, inplace=True)

    # Build custom headers for TEOTIL
    # 1st row
    row1_txt = ('!Dette er en automatisk generert fil. '
                'År:%sVannføring i m3/s. Dato:%s' % (year, pd.to_datetime('now')))
    row1 = [row1_txt] + (len(q_df.columns) - 1)*['',]

    # 2nd row
    row2 = ['Vassdrag', 'QAr', 'Qsnitt', 'Q1', 'Q2', 'Q3', 'Q4', 
            'Q5', 'Q6', 'Q7', 'Q8', 'Q9', 'Q10', 'Q11', 'Q12']

    # Assign header as multi-index
    q_df.columns = pd.MultiIndex.from_tuples(zip(row1, row2))

    # Write output
    out_path = os.path.join(out_fold, 'Q_Vassdrag.txt')
    q_df.to_csv(out_path, sep=';', index=False, encoding='utf-8')
    
def write_teotil_observed_input_file(out_fold, loads_path, engine):
    """ Writes a text file for input to TEOTIL based on observed loads.
    
    Args:
        out_fold:   Str. Path to folder for output .txt file
        loads_path: Str. CSV file of loads, as produced by
                    rid.estimate_loads()
        engine:     SQL-Alchemy 'engine' object already connected to 
                    RESA2
    
    Returns:
        None. Two (identical) text files are written to the specified 
        folder.
    """
    import pandas as pd
    import numpy as np
    import os

    # Read loads data
    lds_df = pd.read_csv(loads_path)
    lds_df.index = lds_df['station_id']

    # Ignore '_Est' cols
    val_cols = [i for i in lds_df.columns if i.split('_')[1]!='Est']    
    lds_df = lds_df[val_cols]

    # Get station data
    sql = ("SELECT station_id, station_name, nve_vassdr_nr AS regine "
           "FROM resa2.stations "
           "WHERE station_id IN %s" % str(tuple(lds_df['station_id'].values)))

    stn_df = pd.read_sql_query(sql, engine)
    stn_df['station_name'] = stn_df['station_name'].str.decode('windows-1252')
    stn_df.index = stn_df['station_id']

    # Join
    lds_df = pd.concat([stn_df, lds_df], axis=1)

    # Tidy
    del lds_df['station_id']
    lds_df.reset_index(inplace=True)

    # Add blank cols for HCHG and SUMPCB
    lds_df['HCHG_tonnes'] = np.nan
    lds_df['SUMPCB_tonnes'] = np.nan

    # Convert Hg to tonnes
    lds_df['Hg_tonnes'] = lds_df['Hg_kg'] / 1000.

    # Reorder
    cols = ['station_id', 'station_name', 'regine', 'SiO2_tonnes', 
            'Cd_tonnes', 'Hg_tonnes', 'Cu_tonnes', 'Pb_tonnes', 'Zn_tonnes', 
            'HCHG_tonnes', 'SUMPCB_tonnes', 'NH4-N_tonnes', 'NO3-N_tonnes', 
            'PO4-P_tonnes', 'TOTN_tonnes', 'TOTP_tonnes', 'SPM_tonnes', 
            'As_tonnes', 'Ni_tonnes', 'TOC_tonnes', 'Cr_tonnes']

    lds_df = lds_df[cols]

    # Build custom headers for TEOTIL
    # 1st row
    row1_txt = ('!Dette er en automatisk generert fil. '
                'Målt vannkjemi. %s' % pd.to_datetime('now'))
    row1 = [row1_txt] + (len(lds_df.columns) - 1)*['',]

    # 2nd row
    row2 = ['!STATION_ID', 'STATION_NAME', 'REGINE', 'SiO2(tonn)', 'Cd(tonn)', 
            'Hg(tonn)', 'Cu(tonn)', 'Pb(tonn)', 'Zn(tonn)', 'HCHG(tonn)', 
            'SUMPCB(tonn)', 'NH4-N(tonn)', 'NO3-N(tonn)', 'PO4-P(tonn)', 
            'TOT-N(tonn)', 'TOT-P(tonn)', 'SPM(tonn)', 'As(tonn)', 'Ni(tonn)', 
            'TOC(tonn)', 'Cr(tonn)']

    # 3rd row
    row3 = ['STATION_ID', 'STATION_NAME', 'REGINE', 'SiO2', 'Cd', 'Hg', 'Cu', 
            'Pb', 'Zn', 'HCHG', 'SUMPCB', 'NH4-N', 'NO3-N', 'PO4-P', 'TOT-N', 
            'TOT-P', 'SPM', 'As', 'Ni', 'TOC', 'Cr']

    # Assign header as multi-index
    lds_df.columns = pd.MultiIndex.from_tuples(zip(row1, row2, row3))

    # Write output (2 identical files with different names)
    # Lower
    out_path = os.path.join(out_fold, 'RID_Observert_Lower.txt')
    lds_df.to_csv(out_path, sep=';', index=False, encoding='utf-8')
    
    # Upper
    out_path = os.path.join(out_fold, 'RID_Observert_Upper.txt')
    lds_df.to_csv(out_path, sep=';', index=False, encoding='utf-8')
    
def get_prev_bio(row, df):
    """ Used for calculating fish farm productivity.
        Returns biomass for previous month. 
        If month = 1, or if data for the previous month
        are not available, returns 0.
        
        df is the same dataframe to which the function is 
        being applied.
    """
    import pandas as pd
    
    # Get row props from multi-index
    loc = row.name[0]
    mon = row.name[1]
    spec = row.name[2]

    # Get value for previous month
    if mon == 1:
        return 0
    else:
        try:
            # Returns a KeyError if data for (mon - 1)
            # do not exist
            return df.loc[(loc, mon - 1, spec)]['biomass']
        
        except KeyError:
            return 0
        
def calc_productivity(row):
    """ Used for calculating fish farm productivity.
        Calculate productivity based on changes in biomass. 
        Duplicated from formula in col H of N_P_ørret_2015.xlsx
    """
    if ((row['biomass'] == 0) or
        (row['biomass_prev'] == 0) or 
        (row['biomass'] < row['biomass_prev'])):
        return row['FORFORBRUK_KILO'] / 1.15
    
    else:
        return row['biomass'] - row['biomass_prev']
    
def calc_tap(row, par):
    """ Used for calculating fish farm productivity.
        Calculate N tap or P tap. I don't understand these calculations,
        but the formulae are duplicated from cols I and J of 
        N_P_ørret_2015.xlsx
    
    Args:
        par: Str. Either 'N' or 'P'.
    """  
    # Define constants
    if par == 'N':
        k1 = 0.0584
        k2 = 0.0296
    elif par == 'P':
        k1 = 0.0096
        k2 = 0.0045
    else:
        raise ValueError('par must be "N" or "P".')
    
    if row['FORFORBRUK_KILO'] == 0:
        return 0
    
    elif row['prod'] == 0:
        return (k1*row['FORFORBRUK_KILO']) - (k2*row['FORFORBRUK_KILO']/1.15)
    
    elif ((k1*row['FORFORBRUK_KILO']) - (k2*row['prod'])) < 0:
        return (k1*row['FORFORBRUK_KILO']) - (k2*row['FORFORBRUK_KILO']/1.15)
    
    else:
        return (k1*row['FORFORBRUK_KILO']) - (k2*row['prod'])        
    
def estimate_fish_farm_nutrient_inputs(fish_df, year, cu):
    """ Main function for estimating fish farm nutrient inputs.
        Recreates the Excel/Access workflow described in 
        section 3.2 of prepare_teotil_inputs.ipynb.
        
        NOTE: Does not yet include calculations for Cu.
        
    Args:
        fish_df: Dataframe. Tidied version of the raw fish farm 
                 data. See fiske_oppdret_2015_raw.xlsx as an 
                 example. Blank rows should be dropped in advance
        year:    Int. Year of interest
        cu:      Int. Annual tonnes of Cu. Distributed in proportion
                 to P loads
    
    Retruns:
        Dataframe in correct format for upload to 
        RESA2.RID_KILDER_AQKULT_VALUES. 
    """
    import pandas as pd
    import numpy as np

    # Fill NaN with 0 where necessary
    cols = ['FISKEBEHOLDNING_ANTALL', 'FISKEBEHOLDNING_SNITTVEKT',
            'TAP_DOD', 'TAP_UTKAST', 'TAP_ROMT', 'TAP_ANNET', 
            'TELLEFEIL', 'UTTAK_KILO']
    for col in cols:
        fish_df[col].fillna(value=0, inplace=True)

    # Calculate biomass    
    fish_df['biomass'] = ((((fish_df['FISKEBEHOLDNING_ANTALL'] * fish_df['FISKEBEHOLDNING_SNITTVEKT']) +
                            (fish_df['TAP_DOD'] * fish_df['FISKEBEHOLDNING_SNITTVEKT']) +
                            (fish_df['TAP_UTKAST'] * fish_df['FISKEBEHOLDNING_SNITTVEKT']) +
                            (fish_df['TAP_ROMT'] * fish_df['FISKEBEHOLDNING_SNITTVEKT']) +
                            (fish_df['TAP_ANNET'] * fish_df['FISKEBEHOLDNING_SNITTVEKT']) +
                            (fish_df['TELLEFEIL'] * fish_df['FISKEBEHOLDNING_SNITTVEKT'])) / 1000.) +
                           fish_df['UTTAK_KILO'])

    # Aggregate by month, location and species
    fish_grp = fish_df.groupby(by=['LOKNR', 'MAANED', 'FISKEARTID'])
    fish_sum = fish_grp.sum()[['FORFORBRUK_KILO', 'biomass']]

    # Get biomass for previous month
    fish_sum['biomass_prev'] = fish_sum.apply(get_prev_bio, args=(fish_sum,), axis=1)

    # Get productivity for each month
    fish_sum['prod'] = fish_sum.apply(calc_productivity, axis=1)

    # Calculate NTAP and PTAP
    fish_sum['ntap'] = fish_sum.apply(calc_tap, args=('N',), axis=1)
    fish_sum['ptap'] = fish_sum.apply(calc_tap, args=('P',), axis=1)

    # Get just the data for trout and salmon
    fish_sum = fish_sum.iloc[(fish_sum.index.get_level_values('FISKEARTID') == 71101) |
                             (fish_sum.index.get_level_values('FISKEARTID') == 71401)]

    # Aggregate by location
    fish_sum.reset_index(inplace=True)
    fish_grp = fish_sum.groupby(by=['LOKNR'])
    fish_sum = fish_grp.sum()[['ntap', 'ptap']]
    
    # Distribute Cu according to P production
    fish_sum['cutap'] = cu*fish_sum['ptap'] / fish_sum['ptap'].sum()

    # Convert to par_ids and melt to format required by RESA2
    fish_sum.columns = [39, 40, 41]
    fish_sum.reset_index(inplace=True)
    fish_sum = pd.melt(fish_sum, id_vars='LOKNR', 
                       var_name='INP_PAR_ID', value_name='VALUE')

    # Add cols to match RESA2 schema
    fish_sum['AR'] = year
    fish_sum['MANED'] = 6
    fish_sum['ART'] = np.nan

    # Reorder cols
    fish_sum = fish_sum[['LOKNR', 'INP_PAR_ID', 'AR', 'MANED', 'ART', 'VALUE']]

    # Rename cols
    fish_sum.columns = ['ANLEGG_NR', 'INP_PAR_ID', 'AR', 'MANED', 'ART', 'VALUE']

    return fish_sum

def estimate_teotil_land_use_coefficients(lu_path, sheetname, 
                                          year, out_fold):
    """ Caclulates land use coefficients for TEOTIL based on data
        supplied by Bioforsk.
        
    Args:
        lu_path:   Str. Path to tidied Excel file from Bioforsk. See
                   jordbruk_2015.xlsx as an example
        sheetname: Str. Name of sheet with data in Bioforsk Excel file
        year:      Int. Year of interest
        out_fold:  Str. Folder in which to save 
                   Koeffisienter_jordbruk.txt
    
    Returns:
        None. The file is written to the specified folder.
    """
    import pandas as pd
    import os
    
    # Read LU areas (same values used every year)
    in_xlsx = (r'C:\Data\James_Work\Staff\Oeyvind_K\Elveovervakingsprogrammet'
               r'\Data\RID_Fylke-Sone_LU_Areas.xlsx')
    lu_areas = pd.read_excel(in_xlsx, sheetname='Sheet1')

    # Read Bioforsk data
    lu_lds = pd.read_excel(lu_path, sheetname=sheetname)

    # Join
    lu_df = pd.merge(lu_lds, lu_areas, how='outer',
                     on='omrade')

    # Calculate required columns
    # N
    lu_df['Naker'] = (lu_df['N_diff']/lu_df['Adyrket_km2'] +
                      lu_df['N_back']/lu_df['Adyrket_km2'])

    lu_df['Npunkt'] = lu_df['N_point']/lu_df['Aeng_km2']

    lu_df['Ndyrket_nat'] = lu_df['N_back']/lu_df['Adyrket_km2']

    # P
    lu_df['Paker'] = (lu_df['P_diff']/lu_df['Adyrket_km2'] +
                      lu_df['P_back']/lu_df['Adyrket_km2'])

    lu_df['Ppunkt'] = lu_df['P_point']/lu_df['Aeng_km2']

    lu_df['Pdyrket_nat'] = lu_df['P_back']/lu_df['Adyrket_km2']

    # Convert areas to da
    lu_df['Adyrket'] = 1000.*lu_df['Adyrket_km2']
    lu_df['Aeng'] = 1000.*lu_df['Aeng_km2']

    # Get cols of interest
    cols = ['Fylke_sone', 'Fysone_navn', 'Naker', 'Npunkt', 'Ndyrket_nat', 
            'Paker', 'Ppunkt', 'Pdyrket_nat', 'Adyrket', 'Aeng']
    lu_df = lu_df[cols]

    # Create multi-index with correct headings for TEOTIL
    first = ['!%s' % year, 'Fysone_navn', 'kg/km2', 'kg/km2', 
             'kg/km2', 'kg/km2', 'kg/km2', 'kg/km2', 'da', 'da']

    secnd = ['Fylke_sone', '', 'Naker', 'Npunkt', 'Ndyrket_nat', 
             'Paker', 'Ppunkt', 'Pdyrket_nat', 'Adyrket', 'Aeng']

    lu_df.columns = pd.MultiIndex.from_tuples(zip(first, secnd))

    # Write output
    out_path = os.path.join(out_fold, 'Koeffisienter_jordbruk.txt')
    lu_df.to_csv(out_path, sep=';', index=False, encoding='utf-8')
    
def get_flow_volumes(stn_df, st_yr, end_yr, engine):
    """ Gets flow summaries for the specified sites. Returns dataframe
        of average annual daily flow volume in 1000m3/day for each
        year between st_yr and end_yr.
    
    Args:
        stn_df: Dataframe sites of interest. Must include 
                [station_id, station_code, station_name]
        st_yr:  Int. Start year of interest
        end_yr: Int. End year of interest
        engine: SQL-Alchemy 'engine' object already connected to RESA2
        
    Returns:
        Dataframe.
    """
    import pandas as pd
    
    # Container for results
    df_list = []

    # Loop over sites
    for stn_id in stn_df['station_id']:
        # Get catch area for chem station
        sql = ("SELECT catchment_area FROM resa2.stations "
               "WHERE station_id = %s" % stn_id)
        area_df = pd.read_sql_query(sql, engine)    
        wc_area = area_df['catchment_area'].iloc[0]

        # Get linked discharge station
        sql = ("SELECT * FROM resa2.default_dis_stations "
               "WHERE station_id = %s" % stn_id)
        dis_df = pd.read_sql_query(sql, engine)
        dis_stn_id = dis_df['dis_station_id'].iloc[0]

        # Get catchment area for discharge station
        sql = ("SELECT area FROM resa2.discharge_stations "
               "WHERE dis_station_id = %s" % dis_stn_id)
        area_df = pd.read_sql_query(sql, engine)    
        dis_area = area_df['area'].iloc[0]

        # Get annual summary flow stats for this station
        sql = ("SELECT TO_CHAR(xdate, 'YYYY') as year, "
               "       AVG(xvalue) as mean, "
               "       MIN(xvalue) as min, " 
               "       MAX(xvalue) as max " 
               "FROM resa2.discharge_values "
               "WHERE dis_station_id = %s "
               "AND xdate >= date '%s-01-01' "
               "AND xdate <= date '%s-12-31' "
               "GROUP BY TO_CHAR(xdate, 'YYYY') "
               "ORDER BY year" % (dis_stn_id, st_yr, end_yr))
        q_df = pd.read_sql_query(sql, engine) 

        # Set index
        q_df.index = q_df['year']
        del q_df['year']

        # Scale flows by area ratio
        q_df = q_df*wc_area/dis_area

        # Convert m3/s to 1000 m3/d
        q_df = q_df*60*60*24/1000

        # Reset index
        q_df.reset_index(inplace=True)

        # Add stn id
        q_df['station_id'] = stn_id

        # Re-order cols to match template
        q_df = q_df[['station_id', 'year', 'mean']]
        q_df.columns = ['station_id', 'year', 'mean_q_1000m3/day']

        # Add to results
        df_list.append(q_df)

    # Combine to single df
    q_df = pd.concat(df_list, axis=0) 

    # Convert year to int
    q_df['year'] = q_df['year'].astype(int)

    return q_df

def spatial_overlays(df1, df2, how='intersection'):
    """Hugely improves performance compared to gpd.overlay(df1, df2, how='intersection').
       An improved version should eventually be available within Geopandas itself.
       
       From here:
           https://github.com/geopandas/geopandas/issues/400
    """
    import geopandas as gpd
    import pandas as pd
    
    df1 = df1.copy()
    df2 = df2.copy()
    df1['geometry'] = df1.geometry.buffer(0)
    df2['geometry'] = df2.geometry.buffer(0)
    if how=='intersection':
        # Spatial Index to create intersections
        spatial_index = df2.sindex
        df1['bbox'] = df1.geometry.apply(lambda x: x.bounds)
        df1['histreg']=df1.bbox.apply(lambda x:list(spatial_index.intersection(x)))
        pairs = df1['histreg'].to_dict()
        nei = []
        for i,j in pairs.items():
            for k in j:
                nei.append([i,k])
        
        pairs = gpd.GeoDataFrame(nei, columns=['idx1','idx2'], crs=df1.crs)
        pairs = pairs.merge(df1, left_on='idx1', right_index=True)
        pairs = pairs.merge(df2, left_on='idx2', right_index=True, 
                            suffixes=['_1','_2'])
        pairs['Intersection'] = pairs.apply(lambda x: 
                                            (x['geometry_1'].intersection(x['geometry_2'])).buffer(0), 
                                            axis=1)
        pairs = gpd.GeoDataFrame(pairs, columns=pairs.columns, crs=df1.crs)
        cols = pairs.columns.tolist()
        cols.remove('geometry_1')
        cols.remove('geometry_2')
        cols.remove('histreg')
        cols.remove('bbox')
        cols.remove('Intersection')
        dfinter = pairs[cols+['Intersection']].copy()
        dfinter.rename(columns={'Intersection':'geometry'}, inplace=True)
        dfinter = gpd.GeoDataFrame(dfinter, columns=dfinter.columns, crs=pairs.crs)
        dfinter = dfinter.loc[dfinter.geometry.is_empty==False]
        return dfinter
    
    elif how=='difference':
        spatial_index = df2.sindex
        df1['bbox'] = df1.geometry.apply(lambda x: x.bounds)
        df1['histreg']=df1.bbox.apply(lambda x:list(spatial_index.intersection(x)))
        df1['new_g'] = df1.apply(lambda x: reduce(lambda x, y: x.difference(y).buffer(0), 
                                                  [x.geometry]+list(df2.iloc[x.histreg].geometry)), 
                                 axis=1)
        df1.geometry = df1.new_g
        df1 = df1.loc[df1.geometry.is_empty==False].copy()
        df1.drop(['bbox', 'histreg', new_g], axis=1, inplace=True)
        return df1